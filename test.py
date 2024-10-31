import json

from api.article_writing import get_summary, list_to_query, get_keywords, extract_content_from_json, shanhuyun_get_body

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tqdm import tqdm


def get_article(article_base, article_choices):
    try:
        doc = Document()
        ref_file = []

        # 设置全局样式
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)

        # 标题
        title_text = article_base['标题']
        title_paragraph = doc.add_paragraph(title_text)
        title_paragraph.style = doc.styles['Heading 1']
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 摘要
        summary = get_summary(article_base, list_to_query(article_choices))
        summary_text = ''.join([chunk.choices[0].delta.content for chunk in summary["ai_say"]])
        summary_paragraph = doc.add_paragraph(summary_text)
        summary_paragraph.style = doc.styles['Body Text']

        # 关键字
        keywords = get_keywords(article_base)
        keywords_text = f"关键词: {keywords}"
        keywords_paragraph = doc.add_paragraph(keywords_text)
        keywords_paragraph.style = doc.styles['Body Text']

        # 正文
        json_list = extract_content_from_json(article_base)
        for index, item in tqdm(enumerate(json_list)):
            print(item)
            if '标题' in item and index != len(json_list) - 1:
                section_title_text = item['标题']
                section_title_paragraph = doc.add_paragraph(section_title_text)
                section_title_paragraph.style = doc.styles['Heading 2']
            elif '小节标题' in item:
                subsection_title_text = item['小节标题']
                subsection_title_paragraph = doc.add_paragraph(subsection_title_text)
                subsection_title_paragraph.style = doc.styles['Heading 3']
                body = shanhuyun_get_body(article_base, str(item), list_to_query(article_choices))
                if body['ref_file']:
                    ref_file.extend(body['ref_file'])
                body_text = ''.join([chunk.choices[0].delta.content for chunk in body['ai_say']])
                body_paragraph = doc.add_paragraph(body_text)
                body_paragraph.style = doc.styles['Body Text']
            elif '标题' in item and index == len(json_list) - 1:
                section_title_text = item['标题']
                section_title_paragraph = doc.add_paragraph(section_title_text)
                section_title_paragraph.style = doc.styles['Heading 2']
                body = shanhuyun_get_body(article_base, str(item), list_to_query(article_choices))
                body_text = ''.join([chunk.choices[0].delta.content for chunk in body['ai_say']])
                body_paragraph = doc.add_paragraph(body_text)
                body_paragraph.style = doc.styles['Body Text']
                if body['ref_file']:
                    ref_file.extend(body['ref_file'])

        # 引用
        if ref_file:
            references_title_text = "引用"
            references_title_paragraph = doc.add_paragraph(references_title_text)
            references_title_paragraph.style = doc.styles['Heading 2']
            for ref in ref_file:
                ref_text = ref
                ref_paragraph = doc.add_paragraph(ref_text)
                ref_paragraph.style = doc.styles['Body Text']

        # 保存文档
        doc.save("output.docx")

    except Exception as e:
        print(e)


if __name__ == '__main__':
    base = '''
    {
"标题": "冰区船舶跟航风险评估及跟航决策整体规划",
"摘要": "本文针对冰区船舶航行中的风险评估及跟航决策问题，结合冰区航行特性、冰级规范、碎冰影响、破冰技术、跟航策略、编队航行、冰阻力以及流体力学等要素，运用CFD-DEM方法，对冰区船舶航行进行深入研究，旨在为冰区船舶的安全航行提供理论依据和实践指导。",
"关键词": "冰区航行，冰级规范，碎冰 brash ice，破冰，跟航，策略，编队，冰阻力 ice resistance，流体力学，CFD-DEM",
"正文": [
{
"标题": "第一章 研究背景及现状",
"内容": "本章主要介绍冰区船舶航行的背景、意义以及国内外研究现状，分析冰区航行中的主要风险因素，为后续研究奠定基础。",烟雨江南，是指江南地区在春雨绵绵时分的独特景致。细雨蒙蒙，如烟似雾，笼罩着水乡的河流、古桥、乌篷船和青石板路。柳树垂下柔软的枝条
"小节": [
{
"小节标题": "1.1 冰区航行的背景与意义",
"内容": "阐述冰区航行的重要性以及对于全球航运的影响。"
},
{
"小节标题": "1.2 国内外研究现状",
"内容": "分析国内外在冰区船舶航行领域的研究进展和成果。"
},
{
"小节标题": "1.3 冰区航行风险因素分析",
"内容": "探讨冰区航行中可能遇到的风险因素及其影响。"
},
{
"小节标题": "1.4 论文结构安排",
"内容": "介绍本论文的研究内容和结构安排。"
}
]
},
{
"标题": "第二章 冰区船舶航行理论基础",
"内容": "本章介绍冰区船舶航行的相关理论基础，包括冰级规范、冰力学特性、流体力学原理等。",
"小节": [
{
"小节标题": "2.1 冰级规范与船舶设计",
"内容": "解析冰级规范对船舶设计的影响和要求。"
},
{
"小节标题": "2.2 海冰的物理力学特性",
"内容": "分析海冰的类型、强度、速度等物理力学特性。"
},
{
"小节标题": "2.3 流体力学在冰区航行中的应用",
"内容": "探讨流体力学原理在冰区船舶航行中的应用。"
},
{
"小节标题": "2.4 冰区船舶阻力分析",
"内容": "研究冰区船舶航行时的阻力特性和影响因素。"
}
]
},
{
"标题": "第三章 冰区船舶跟航风险评估",
"内容": "本章对冰区船舶跟航过程中的风险进行评估，包括风险识别、风险评估和风险控制。",
"小节": [
{
"小节标题": "3.1 风险识别",
"内容": "识别冰区船舶跟航过程中的潜在风险因素。"
},
{
"小节标题": "3.2 风险评估方法",
"内容": "介绍冰区船舶跟航风险评估的方法和工具。"
},
{
"小节标题": "3.3 风险控制策略",
"内容": "提出针对冰区船舶跟航风险的预防和控制措施。"
},
{
"小节标题": "3.4 案例分析",
"内容": "通过具体案例对冰区船舶跟航风险评估进行实证分析。"
}
]
},
{
"标题": "第四章 冰区船舶跟航决策模型构建",
"内容": "本章构建冰区船舶跟航决策模型，包括决策变量、目标函数和约束条件等。",
"小节": [
{
"小节标题": "4.1 决策变量与参数设定",
"内容": "定义冰区船舶跟航决策模型中的决策变量和参数。"
},
{
"小节标题": "4.2 目标函数构建",
"内容": "构建冰区船舶跟航决策模型的目标函数。"
},
{
"小节标题": "4.3 约束条件分析",
"内容": "分析冰区船舶跟航决策模型的约束条件。"
},
{
"小节标题": "4.4 模型求解方法",
"内容": "探讨冰区船舶跟航决策模型的求解方法和算法。"
}
]
},
{
"标题": "第五章 冰区船舶跟航决策实证研究",
"内容": "本章通过实证研究，验证冰区船舶跟航决策模型的有效性和实用性。",
"小节": [
{
"小节标题": "5.1 数据收集与处理",
"内容": "介绍实证研究所需数据的收集和处理方法。"
},
{
"小节标题": "5.2 模型应用与结果分析",
"内容": "将构建的决策模型应用于实际案例，并分析结果。"
},
{
"小节标题": "5.3 结果讨论与启示",
"内容": "对实证研究结果进行讨论，并提出相应的管理启示。"
},
{
"小节标题": "5.4 模型优化建议",
"内容": "根据实证研究结果，提出冰区船舶跟航决策模型的优化建议。"
}
]
},
{
"标题": "第六章 冰区船舶编队航行研究",
"内容": "本章研究冰区船舶编队航行的策略和方法，以提高航行效率和安全性。",
"小节": [
{
"小节标题": "6.1 编队航行原理",
"内容": "介绍冰区船舶编队航行的基本原理和优势。"
},
{
"小节标题": "6.2 编队航行策略",
"内容": "探讨冰区船舶编队航行的策略选择和实施。"
},
{
"小节标题": "6.3 编队航行风险评估",
"内容": "分析冰区船舶编队航行中的风险因素和评估方法。"
},
{
"小节标题": "6.4 编队航行实证研究",
"内容": "通过实证研究验证冰区船舶编队航行策略的有效性。"
}
]
},
{
"标题": "第七章 冰区船舶跟航决策支持系统开发",
"内容": "本章开发冰区船舶跟航决策支持系统，以提高冰区航行的决策效率。",
"小节": [
{
"小节标题": "7.1 系统需求分析",
"内容": "分析冰区船舶跟航决策支持系统的功能和性能需求。"
},
{
"小节标题": "7.2 系统设计与开发",
"内容": "介绍冰区船舶跟航决策支持系统的设计与开发过程。"
},
{
"小节标题": "7.3 系统测试与评估",
"内容": "对开发的决策支持系统进行测试和性能评估。"
},
{
"小节标题": "7.4 系统应用案例",
"内容": "展示冰区船舶跟航决策支持系统的实际应用案例。"
}
]
},
{
"标题": "第八章 发展趋势与展望",
"内容": "本章展望冰区船舶航行的发展趋势，提出未来研究方向和策略。",
"小节": [
{
"小节标题": "8.1 冰区船舶航行发展趋势",
"内容": "分析冰区船舶航行的未来发展趋势。"
},
{
"小节标题": "8.2 冰区船舶跟航决策研究展望",
"内容": "展望冰区船舶跟航决策研究的未来方向。"
},
{
"小节标题": "8.3 研究局限与未来工作",
"内容": "讨论本研究的局限性，并提出未来的研究方向。"
}
]
},
{
"标题": "第九章 结论与建议",
"内容": "本章总结本研究的主要成果，并提出相应的建议。",
"小节": [
{
"小节标题": "9.1 研究结论",
"内容": "概括本研究的主要结论和发现。"
},
{
"小节标题": "9.2 实践建议",
"内容": "提出基于研究结果的管理和实践建议。"
},
{
"小节标题": "9.3 研究贡献",
"内容": "阐述本研究在理论和实践中的贡献。"
},
{
"小节标题": "9.4 研究展望",
"内容": "展望未来冰区船舶跟航决策研究的可能方向。"
}
]
}
]
}
    '''
    base = json.loads(base)
    print(base)
    get_article(base, ["jianting"])
