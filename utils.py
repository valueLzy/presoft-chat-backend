import hashlib
import os
import uuid
from llama_index.core import Settings
from llama_index.core import SimpleDirectoryReader, StorageContext, KnowledgeGraphIndex
from llama_index.core.graph_stores import SimpleGraphStore

from llm.glm4_llamaindex import GLM4
from llm.llama_index_embeddings import InstructorEmbeddings
import requests
from docx import Document
from docx.shared import RGBColor
from minio import Minio
from knowledge.dataset_api import matching_paragraph
from llm.embeddings import  rerank

minio_client = Minio(
    "192.168.2.8:19000",
    access_key="minioadmin",
    secret_key="minioadmin",
    secure=False,
)


def md5_encrypt(password):
    # 创建一个md5哈希对象
    md5 = hashlib.md5()

    # 更新哈希对象以包含密码的二进制数据
    md5.update(password.encode('utf-8'))

    # 获取加密后的十六进制表示
    encrypted_password = md5.hexdigest()

    return encrypted_password


def download_file(bucket_name: str, file_name: str):
    unique_id = str(uuid.uuid4())
    folder_path = f'./data/{unique_id}'
    os.makedirs(folder_path, exist_ok=True)
    file_path = os.path.join(folder_path, file_name)
    try:
        minio_client.fget_object(
            bucket_name=bucket_name,
            object_name=file_name,
            file_path=file_path,
        )
        return {
            "file_path": str(file_path),
            "file_dir": unique_id
        }
    except Exception as e:
        return ""


def put_file(bucket_name: str, file_name: str, file_path: str) -> bool:
    try:
        minio_client.fput_object(
            bucket_name=bucket_name,
            object_name=file_name,
            file_path=file_path,
        )
        return True
    except Exception as e:
        return False


def has_japanese(text: str) -> bool:
    for char in str(text):
        if ('\u3040' <= char <= '\u30FF') or ('\u31F0' <= char <= '\u31FF'):
            return True
    return False


def get_red_text_from_docx(file_path):
    # 打开文档
    doc = Document(file_path)
    red_texts = []

    # 遍历所有段落
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # 检查字体颜色是否为红色
            if run.font.color and run.font.color.rgb == RGBColor(255, 0, 0):
                red_texts.append(run.text)

    # 遍历所有表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        # 检查字体颜色是否为红色
                        if run.font.color and run.font.color.rgb == RGBColor(255, 0, 0):
                            red_texts.append(run.text)

    return red_texts


def replace_text_in_docx(file_path, replacements, new_file_path):
    # 打开文档
    doc = Document(file_path)

    # 遍历所有段落并替换文本
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                replace_text_in_paragraph(paragraph, old_text, new_text)

    # 遍历所有表格并替换文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in paragraph.text:
                            replace_text_in_paragraph(paragraph, old_text, new_text)

    # 保存修改后的文档
    doc.save(new_file_path)


def replace_text_in_paragraph(paragraph, old_text, new_text):
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            run.font.color.rgb = RGBColor(72, 116, 203)


def parse_file_other(bucket_name: str, file_name: str) -> list:
    # 定义请求体
    url = "http://192.168.2.8:5050/knowledge/parse-other/"
    data = {
        "bucket_name": bucket_name,  # 替换为你使用的模型名
        "object_name": file_name
    }

    # 发起POST请求
    response = requests.post(url, json=data)
    # 处理响应
    if response.status_code == 200:
        result = response.json()
        return result
    else:
        return []


def parse_file_pdf(bucket_name: str, file_name: str) -> list:
    # 定义请求体
    url = "http://192.168.2.8:5050/knowledge/parse-pdf/"
    data = {
        "bucket_name": bucket_name,  # 替换为你使用的模型名
        "object_name": file_name
    }

    # 发起POST请求
    response = requests.post(url, json=data)
    # 处理响应
    if response.status_code == 200:
        result = response.json()
        return result
    else:
        return []


def matching_milvus_paragraph(query, collection_name, matches_number):
    ref = matching_paragraph(query, collection_name, 1000)
    filtered_results = []
    for result in ref:
        filtered_result = [res.entity.text for res in result if res.score > 0]
        filtered_results.append(filtered_result)
    a = rerank(query, filtered_results[0], matches_number)
    rerank_results = [y['index'] for y in a if y['relevance_score'] > 0.7]
    rerank_filtered_result_text = []
    for index in rerank_results:
        rerank_filtered_result_text.append(ref[0][index].fields['text'] + '\n')
    return rerank_filtered_result_text


def convert_to_vis_format(data):
    nodes = set()
    edges = []
    for from_node, relationships in data.items():
        nodes.add(from_node)
        for relationship in relationships:
            label, to_node = relationship
            nodes.add(to_node)
            edges.append({
                'from': from_node,
                'to': to_node,
                'label': label
            })
    return edges


def get_graph(bucket_name, file_name):
    download_file_res = download_file(bucket_name, file_name)
    file_dir = download_file_res['file_dir']

    Settings.embed_model = InstructorEmbeddings()
    Settings.llm = GLM4()

    documents = SimpleDirectoryReader(f"./data/{file_dir}/").load_data()

    graph_store = SimpleGraphStore()

    storage_context = StorageContext.from_defaults(graph_store=graph_store)

    index = KnowledgeGraphIndex.from_documents(documents=documents,
                                               max_triplets_per_chunk=3,
                                               storage_context=storage_context,
                                               include_embeddings=True)
    g = index.graph_store._data.graph_dict
    return str(convert_to_vis_format(g)).replace("'", "\"")


if __name__ == '__main__':
    # print(put_file("vue-file", "aaaa.xlsx", "data/aaaa.xlsx"))
    print(get_graph("vue-file", "入职指南.pdf"))
