import json
import os

import openpyxl
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import RGBColor
from starlette.websockets import WebSocket

from api.ordinary_conversations import ordinary
from llm.glm4 import glm4_9b_chat_ws, glm4_9b_chat_http
from util.websocket_utils import ConnectionManager


def has_japanese(text):
    for char in str(text):
        if ('\u3040' <= char <= '\u30FF') or ('\u31F0' <= char <= '\u31FF'):
            return True
    return False


# 普通LLM
def del_ordinary(temperature, prompt, historical_dialogue):
    return ordinary(temperature, prompt, historical_dialogue)


# 日语修正助手 ##文件
async def del_japanese_file(websocket,file_path):
    dot_index = file_path.rfind('.')
    file_type = file_path[dot_index + 1:] if dot_index != -1 else ''
    new_file_path = "modified_" + file_path
    # EXCEL
    if file_type == 'xlsx' or file_type == 'xls':
        # 加载工作簿和工作表
        workbook = openpyxl.load_workbook(file_path)
        for sheet in workbook.sheetnames:
            worksheet = workbook[sheet]
            # 遍历每个单元格
            for row in worksheet.iter_rows(min_row=1, max_col=worksheet.max_column, max_row=worksheet.max_row):
                for cell in row:
                    if isinstance(cell.value, str) and has_japanese(cell.value):
                        # 处理日文内容
                        new_value = del_japanese_prompt(0.1, cell.value, [])
                        cell_value = str(cell.value).replace("\n", "")
                        new_cell_value = str(new_value).replace("\n", "")
                        manager = ConnectionManager()
                        await manager.connect(websocket)
                        await manager.send_personal_message(json.dumps({
                            "newpart": {
                                "sheet": sheet,
                                "row": cell.row,
                                "column": get_column_letter(cell.column),
                                "before": cell_value,
                                "after": new_cell_value
                            }
                        }, ensure_ascii=False), websocket)
                        # cell.value = new_value
            # # 如果工作表被处理，保存修改
            # if processed:
            #     workbook.save(new_file_path)

        # # 返回新文件的内容
        # with open(new_file_path, "rb") as file:
        #     file_contents = file.read()
        # os.remove(new_file_path)
        # yield {
        #     "name": new_file_path,
        #     "content": file_contents,
        #     "type": "xlsx"
        # }
    # # TXT
    # if file_type == 'txt':
    #     # 读取文件内容
    #     text = file_path.getvalue().decode("utf-8")
    #     text_array = text.split("\n\n")
    #     for index, item in enumerate(text_array):
    #         if has_japanese(item):
    #             ai_value = del_japanese_prompt(0.1, item, [])
    #             aisay_item = item.replace("\n", "")
    #             aisay_value = ai_value.replace("\n", "")
    #             text_array[index] = ai_value  # 替换原始文本数组中的值
    #             yield f'''
    #                 修正：
    #                 - **修正前**: {aisay_item}
    #                 - **修正後**: {aisay_value}
    #             '''
    #
    #     # 将修改后的文本数组重新组合成字符串
    #     new_text = '\n\n'.join(text_array)
    #
    #     yield {
    #         "name": new_file_path,
    #         "content": new_text,
    #         "type": "txt"
    #     }
    # #DOC
    # if file_type == 'doc' or file_type == 'docx':
    #     replacements = {}
    #     red_list = get_red_text_from_docx(file_path)
    #     if len(red_list) > 0:
    #         for index, item in enumerate(red_list):
    #             ai_value = del_japanese_prompt(0.1, item, [])
    #             aisay_item = item.replace("\n", "")
    #             aisay_value = ai_value.replace("\n", "")
    #             replacements[item] = ai_value
    #             yield f'''
    #                 修正：
    #                 - **修正前**: {aisay_item}
    #                 - **修正後**: {aisay_value}
    #             '''
    #             replace_text_in_docx(file_path, replacements, new_file_path)
    #         # 返回新文件的内容
    #         with open(new_file_path, "rb") as file:
    #             file_contents = file.read()
    #         os.remove(new_file_path)
    #         yield {
    #             "name": new_file_path,
    #             "content": file_contents,
    #             "type": "docx"
    #         }
    #     else:
    #         yield '翻訳が必要なテキストは検出されませんでした。'


# 日语修正助手 ##文本
def del_japanese_prompt(temperature, prompt, history):
    # '私はあなたが日本語の翻訳、校正修辞の改善の役を担当することを望んでいます。\
    # 私はどんな言語であなたと交流することができて、あなたは言語を識別することができて、\
    # それを翻訳してそしてもっと正確で合理的な日本語で私に答えることができます。\
    # 私の標準ではない日本語の文をより標準的な尊敬の表現に修正して、意味が変わらないようにしてください。\
    # 修正した文をそのまま返します！！！！\n'
    system = '''
            '我希望你能担当日语校对改进的角色。\n
             我只用日语和你交流，你要用最准确合理的日语回答我。\n
             把我不标准的日语句子修改成没有拼写和语法错误、地道的标准日语，并且意思不要改变，尽可能和原句保持一致。\n
             如果有数字和符号，请不要修改。\n
             请只返回我修改后的句子，不要其他多余的解释!!!'
    '''
    history.insert(0, {"role": "system", "content": system})
    history.append({"role": "user", "content": prompt})
    return glm4_9b_chat_http(history,temperature)


def get_red_text_from_docx(file_path):
    # 打开文档
    doc = Document(file_path)
    red_texts = []

    # 遍历所有段落
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # 检查字体颜色是否为红色
            if run.font.color and run.font.color.rgb == RGBColor(255, 0, 0):
                red_texts.append(run.text)

    # 遍历所有表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        # 检查字体颜色是否为红色
                        if run.font.color and run.font.color.rgb == RGBColor(255, 0, 0):
                            red_texts.append(run.text)

    return red_texts


def replace_text_in_docx(file_path, replacements, new_file_path):
    # 打开文档
    doc = Document(file_path)

    # 遍历所有段落并替换文本
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                replace_text_in_paragraph(paragraph, old_text, new_text)

    # 遍历所有表格并替换文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old_text, new_text in replacements.items():
                        if old_text in paragraph.text:
                            replace_text_in_paragraph(paragraph, old_text, new_text)

    # 保存修改后的文档
    doc.save(new_file_path)


def replace_text_in_paragraph(paragraph, old_text, new_text):
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            run.font.color.rgb = RGBColor(72, 116, 203)
