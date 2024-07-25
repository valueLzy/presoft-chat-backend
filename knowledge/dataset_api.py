
import tiktoken
import pandas as pd
from docx import Document
import fitz
from langchain_text_splitters import RecursiveCharacterTextSplitter, CharacterTextSplitter

from llm.embeddings import segment_document, bg3_m3, rerank

from milvus.milvus_tools import get_milvus_collections_info, create_milvus, delete_milvus, insert_milvus, query_milvus,search_milvus


def text_splitter(text, length_function, separators, chunk_size, chunk_overlap):
    splits = []
    try:
        splitter = None
        if length_function == 'Character':
            length_function = len
        elif length_function == 'Tokens':
            enc = tiktoken.get_encoding("cl100k_base")

            def length_function(text: str) -> int:
                return len(enc.encode(text))
        if text_splitter == 'RecursiveCharacter':
            if len(separators) == 0:
                separators = ["\n\n", "\n", " ", ""]
            splitter = RecursiveCharacterTextSplitter(
                # ["\n\n", "\n", " ", ""]
                separators=separators,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                length_function=length_function,
            )
        elif text_splitter == 'Character':
            if len(separators) == 0:
                separators = ["\n\n"]
            splitter = CharacterTextSplitter(
                # \n\n
                separator=separators[0],  # Split character (default \n\n)
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                length_function=length_function,
            )
        res = splitter.split_text(text)
        return res
    except Exception as e:
        print(f"方法text_splitter报错: {e}")
        return splits


def get_knowledge_name():
    return get_milvus_collections_info()


def create_milvus_collection(collection_name, description):
    return create_milvus(collection_name, description)


def konwledge_division_type(file_path):
    file_type = str(file_path.name).split('.')[-1]
    if file_type == 'xlsx' or file_type == 'xls':
        return dispose_excel(file_path)
    elif file_type == 'txt':
        return dispose_txt(file_path)
    elif file_type == 'doc' or file_type == 'docx':
        return dispose_word(file_path)
    elif file_type == 'pdf':
        return dispose_pdf(file_path)


def delete_milvus_collection(collection_name):
    return delete_milvus(collection_name)


def dispose_excel(file_path):
    # 读取 Excel 文件的所有表单
    xls = pd.ExcelFile(file_path)

    # 初始化一个列表来存储所有表单的文字
    text_list = []

    # 遍历每一个表单
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(file_path, sheet_name=sheet_name, header=None)

        # 遍历每一行
        for index, row in df.iterrows():
            # 将每一行转换为一个字符串
            text = row.astype(str).values
            row_text = ' '.join(text).replace("nan", "")
            if not row_text.strip():
                print(row_text)
            else:
                text_list.append(row_text)

    return text_list


def dispose_txt(file_path):
    try:
        text = file_path.getvalue().decode("utf-8")
        return segment_document(text)
    except Exception as e:
        return f"读取文件时出错: {e}"


def dispose_word(file_path):
    # 打开文档
    doc = Document(file_path)
    texts = []

    # 遍历所有段落
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if len(run.text.strip()) > 1000:
                texts.append(segment_document(run.text))
            else:
                texts.append(run.text)

    # 遍历所有表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if len(run.text.strip()) > 1000:
                            texts.append(segment_document(run.text))
                        else:
                            texts.append(run.text)

    return texts


def dispose_pdf(file_path):
    text = ""
    text_list = []
    try:
        # 将上传的文件内容读取为字节流
        pdf_bytes = file_path.read()

        # 打开字节流作为一个PyMuPDF文档
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")

        # 提取每一页的文本内容
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text_list.append(page.get_text())
            text += page.get_text()

        # 关闭文档
        doc.close()
        segment = segment_document(text)
    except Exception as e:
        print(f"Error occurred: {e}")
    return text_list if len(segment) == 1 else segment


def batch_insert_knowledge(collection_name, data):
    try:
        for item in data:
            if item is None:
                continue
            data = [{
                'text': item,
                'embeddings': bg3_m3(item)
            }]
            insert_milvus(data, collection_name)
        return 'success'
    except Exception as e:
        return e


def matching_paragraph(text, collection_name, matches_number):
    res = search_milvus(bg3_m3(text), collection_name, matches_number)
    return res


def get_paragraph(collection_name):
    paragraph_list = []
    res = query_milvus(collection_name)
    for item in res:
        paragraph_list.append(item['text'])
    return paragraph_list


if __name__ == '__main__':
    # print(get_paragraph('lzy'))
    query = '''
    板架材料为Q235钢,密度为多少
    '''

    res = matching_paragraph(query, 'damage_explosion_v1', 1000)
    filtered_results = []
    for result in res:
        filtered_result = [res.entity.text for res in result if res.score > 0]
        filtered_results.append(filtered_result)
    a = rerank(query,filtered_results[0])
    rerank_results = [y['index'] for y in a if y['relevance_score'] > 0.4]
    print(rerank_results)